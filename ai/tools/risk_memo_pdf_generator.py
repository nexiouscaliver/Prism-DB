import os
import logging
import jinja2
from weasyprint import HTML, default_url_fetcher
from datetime import datetime
from db.utils import get_all_sections_from_risk_memo, get_field_from_risk_memo
from docx import Document
from docx.shared import RGBColor, Pt
from markdown import markdown as md_to_html
from markdown2 import markdown as md_to_text  # For plain text conversion
from bs4 import BeautifulSoup  # Ensure BeautifulSoup is installed
from dotenv import load_dotenv
import boto3
from botocore.exceptions import NoCredentialsError
import mimetypes

load_dotenv()

# Configure logging for debugging - set level to DEBUG
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class RiskMemoGenerator:
    def __init__(self):
        # Set paths to templates and output folder
        self.templates_dir = os.path.join(os.path.dirname(__file__), "templates")
        self.css_path = os.path.join(self.templates_dir, "rm_styles.css")
        self.template_path = os.path.join(self.templates_dir, "rm_template.html")

        self.env = os.getenv("ENV", "development")
        
        # Setup for different environments
        if self.env == "production":
            # In production, use S3 for static assets 
            self.s3_bucket = os.getenv("S3_BUCKET_NAME")
            self.static_assets_prefix = os.getenv("S3_STATIC_ASSETS_PREFIX", "static")
            self.static_url = f"https://{self.s3_bucket}.s3.amazonaws.com/{self.static_assets_prefix}"
            
            # Output folder for generated files
            self.output_folder = os.getenv("GENERATED_FILES_S3_PATH")
            
            # Initialize S3 client
            self.s3_client = boto3.client(
                "s3",
                aws_access_key_id=os.getenv("AWS_ACCESS_KEY_ID"),
                aws_secret_access_key=os.getenv("AWS_SECRET_ACCESS_KEY"),
                region_name=os.getenv("AWS_REGION"),
            )
            
            # Upload static assets to S3 if needed
            try:
                self.upload_static_assets_to_s3()
            except Exception as e:
                logging.warning(f"Failed to upload static assets to S3: {e}")
        else:
            # In development, use local file paths
            self.static_url = f"file://{self.templates_dir.replace(os.path.sep, '/')}"
            
            # Output folder for generated files
            self.output_folder = os.getenv(
                "GENERATED_FILES_LOCAL_PATH",
                os.path.join(os.path.dirname(__file__), "..", "..", "generated_files"),
            )
            # Ensure the local output directory exists
            os.makedirs(self.output_folder, exist_ok=True)

    def upload_to_s3(self, file_path, file_name):
        try:
            bucket_name = self.output_folder.replace("s3://", "").split("/")[0]
            s3_key = (
                "/".join(self.output_folder.replace("s3://", "").split("/")[1:])
                + f"/{file_name}"
            )
            self.s3_client.upload_file(file_path, bucket_name, s3_key)
            logging.info(
                f"Uploaded {file_name} to S3 bucket '{bucket_name}' at '{s3_key}'."
            )
        except FileNotFoundError:
            logging.error("The file was not found.")
        except NoCredentialsError:
            logging.error("AWS credentials not available.")

    def create_pdf(self, ticker):
        try:
            # Generate current date and time
            current_datetime = datetime.now()
            date = current_datetime.strftime("%d %b %y")
            time = current_datetime.strftime("%H:%M")

            # Retrieve all sections from the risk memo
            sections = get_all_sections_from_risk_memo(ticker)
            if not sections:
                logging.error(f"No sections found for ticker '{ticker}'.")
                return

            # Debug logging to see what we actually have in the database
            logger.debug(f"Retrieved sections for {ticker}: {list(sections.keys())}")

            # Read the HTML template content
            with open(self.template_path, "r") as file:
                template_content = file.read()

            template = jinja2.Template(template_content)

            # Updated mapping that matches actual keys in MongoDB (lowercase, spaces)
            section_mapping = {
                "business": "Business",
                "risks and mitigation": "Risks and Mitigation",
                "market risks": "Market Risks",
                "management perspective": "Management Perspective",
                "legal proceedings": "Legal Proceedings",
                "supplementary financial data": "Supplementary Financial Data",
                "financial statements": "Financial Statements",
                "peer comparison": "Peer Comparison",
                "analysts comments": "Analysts Comments",
                "profile": "Profile",
                "financial_overview": "Financial Overview",
            }

            # Prepare data for the template using explicit section keys
            template_sections = {}
            for db_key, template_key in section_mapping.items():
                content = sections.get(db_key, "")
                logger.debug(f"Section {db_key} -> {template_key}: {'Found' if content else 'MISSING'}")
                if db_key == "financial_overview":
                    # Financial overview is already HTML, don't convert
                    template_sections[template_key] = content
                else:
                    # Other sections need markdown to HTML conversion
                    template_sections[template_key] = md_to_html(content)

            data = {
                "Title": "Credit Risk Memo",
                "Company": get_field_from_risk_memo(ticker, "companyName"),
                "Ticker": ticker.upper(),
                "Date": date,
                "Time": time,
                "static_url": self.static_url,  
                "Sections": template_sections,

            }

            # Render the template with the data
            rm_html = template.render(**data)

            # Generate the PDF
            current_date = current_datetime.strftime("%d%m%y")
            pdf_file_name = f"{ticker}_Risk_Memo_{current_date}.pdf"
            pdf_output_path = os.path.join(self.output_folder, pdf_file_name)

            # Create HTML object
            html = HTML(string=rm_html, url_fetcher=self.custom_url_fetcher)

            # Write PDF
            html.write_pdf(pdf_output_path)

            logging.debug(f"PDF generated and saved to '{pdf_output_path}'.")

            if self.env == "production":
                self.upload_to_s3(pdf_output_path, pdf_file_name)
                os.remove(pdf_output_path)  # Optional: Remove local file after upload

            # Create DOCX and capture the filename
            docx_file_name = self.create_docx(ticker)

            # Return the filename without '^'
            return pdf_file_name

        except Exception as e:
            logging.error(f"An error occurred while creating PDF: {e}")

    def create_docx(self, ticker):
        try:
            # Generate current date and time
            current_datetime = datetime.now()
            date = current_datetime.strftime("%d %b %y")
            time = current_datetime.strftime("%H:%M")

            # Retrieve all sections from the risk memo
            sections = get_all_sections_from_risk_memo(ticker)
            if not sections:
                logging.error(f"No sections found for ticker '{ticker}'.")
                return
                
            # Debug logging to see what sections we have
            logger.debug(f"Retrieved sections for DOCX: {list(sections.keys())}")

            # Updated mapping that matches actual keys in MongoDB (lowercase, spaces)
            section_mapping = {
                "business": "Business",
                "risks and mitigation": "Risks and Mitigation",
                "market risks": "Market Risks",
                "management perspective": "Management Perspective",
                "legal proceedings": "Legal Proceedings",
                "supplementary financial data": "Supplementary Financial Data",
                "financial statements": "Financial Statements",
                "peer comparison": "Peer Comparison",
                "analysts comments": "Analysts Comments",
                "profile": "Profile",
                "financial_overview": "Financial Overview",
            }

            # Prepare data for the template using explicit section keys
            prepared_sections = {}
            for db_key, template_key in section_mapping.items():
                content = sections.get(db_key, "")
                logger.debug(f"DOCX Section {db_key} -> {template_key}: {'Found' if content else 'MISSING'}")
                if db_key == "financial_overview":
                    # Financial overview is already HTML, don't convert
                    prepared_sections[template_key] = content
                else:
                    # Other sections need markdown to HTML conversion
                    prepared_sections[template_key] = md_to_html(content)

            # Create DOCX
            self._create_docx(ticker, prepared_sections, date, time)

        except Exception as e:
            logging.error(f"An error occurred while creating DOCX: {e}")
            logging.error(f"Exception details: {str(e)}")

    def _create_docx(self, ticker, sections, date, time):
        try:
            # Initialize a new Word Document
            document = Document()

            # Add Title
            document.add_heading("Credit Risk Memo", level=0)

            # Add Company Info
            document.add_paragraph(
                f"Company: {get_field_from_risk_memo(ticker, 'companyName')}"
            )
            document.add_paragraph(f"Ticker: {ticker.upper()}")
            document.add_paragraph(f"Date: {date}")
            document.add_paragraph(f"Time: {time}")
            document.add_paragraph("")  # Add an empty paragraph for spacing

            # Add Sections
            for section_title, html_content in sections.items():
                # Add Section Heading
                heading = section_title.replace("_", " ").title()
                logging.debug(f"Adding section heading: {heading}")
                document.add_heading(heading, level=1)

                # Parse HTML and add to document
                self.html_to_docx(document, html_content)

            # Save the Document
            current_date = datetime.now().strftime("%d%m%y")
            docx_file_name = f"{ticker}_Risk_Memo_{current_date}.docx"
            docx_output_path = os.path.join(self.output_folder, docx_file_name)
            document.save(docx_output_path)

            logging.info(f"DOCX generated and saved to '{docx_output_path}'.")

            if self.env == "production":
                self.upload_to_s3(docx_output_path, docx_file_name)
                os.remove(docx_output_path)  # Optional: Remove local file after upload

            return docx_file_name

        except Exception as e:
            logging.error(f"An error occurred while creating DOCX: {e}")

    def html_to_docx(self, document, html_content):
        try:
            soup = BeautifulSoup(html_content, "html.parser")

            # If the HTML content doesn't have a body tag, use the entire soup
            body = soup.body
            if body:
                elements = body.contents
            else:
                elements = soup.contents

            for element in elements:
                self._add_html_element_to_docx(document, element)
        except Exception as e:
            logging.error(f"An error occurred while parsing HTML to DOCX: {e}")

    def _add_html_element_to_docx(self, document, element):
        if isinstance(element, str):
            # Plain text
            if element.strip():
                document.add_paragraph(element)
            return

        if element.name == "h1":
            document.add_heading(element.get_text(), level=1)
        elif element.name == "h2":
            document.add_heading(element.get_text(), level=2)
        elif element.name == "h3":
            document.add_heading(element.get_text(), level=3)
        elif element.name == "p":
            paragraph = document.add_paragraph()
            self._add_html_content_to_paragraph(paragraph, element)
        elif element.name == "ul":
            for li in element.find_all("li", recursive=False):
                paragraph = document.add_paragraph(style="List Bullet")
                self._add_html_content_to_paragraph(paragraph, li)
        elif element.name == "ol":
            for li in element.find_all("li", recursive=False):
                paragraph = document.add_paragraph(style="List Number")
                self._add_html_content_to_paragraph(paragraph, li)
        elif element.name == "blockquote":
            paragraph = document.add_paragraph(style="Intense Quote")
            self._add_html_content_to_paragraph(paragraph, element)
        elif element.name == "pre":
            code_text = element.get_text()
            paragraph = document.add_paragraph()
            run = paragraph.add_run(code_text)
            run.font.name = "Courier New"  # Monospaced font for code
            run.font.size = Pt(10)  # Optional: set font size
        else:
            # Handle other tags if needed
            if element.name:
                logging.debug(f"Unhandled HTML tag in DOCX conversion: {element.name}")

    def _add_html_content_to_paragraph(self, paragraph, element):
        for child in element.children:
            if isinstance(child, str):
                paragraph.add_run(child)
            elif child.name in ["strong", "b"]:
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif child.name in ["em", "i"]:
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif child.name == "u":
                run = paragraph.add_run(child.get_text())
                run.underline = True
            elif child.name == "a":
                run = paragraph.add_run(child.get_text())
                run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Blue color for links
                # Optionally, add hyperlink functionality if needed
            elif child.name == "br":
                paragraph.add_run("\n")
            else:
                # Handle other inline elements if needed
                pass

    def custom_url_fetcher(self, url):
        """Custom URL fetcher that handles both file:// URLs and S3 URLs."""
        logging.debug(f"Fetching URL: {url}")
        
        # If it's an S3 URL and we're in production
        if self.env == "production" and url.startswith(f"https://{self.s3_bucket}.s3.amazonaws.com"):
            try:
                # Parse the URL to get the S3 key
                s3_key = url.replace(f"https://{self.s3_bucket}.s3.amazonaws.com/", "")
                
                # Create a temporary file to store the S3 content
                import tempfile
                with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
                    # Download the file from S3
                    self.s3_client.download_file(self.s3_bucket, s3_key, tmp_file.name)
                    
                    # Read the file content
                    with open(tmp_file.name, 'rb') as f:
                        content = f.read()
                    
                    # Get the content type
                    content_type = mimetypes.guess_type(url)[0] or 'application/octet-stream'
                    
                    # Clean up the temporary file
                    os.unlink(tmp_file.name)
                    
                    # Return the content in the format expected by WeasyPrint
                    return {'string': content, 'mime_type': content_type}
            except Exception as e:
                logging.error(f"Error fetching from S3: {e}")
                # Fall back to default fetcher
                return default_url_fetcher(url)
        
        # For file:// URLs or non-S3 URLs, use the default fetcher
        return default_url_fetcher(url)

    def upload_static_assets_to_s3(self):
        """Upload template static assets (CSS files and images) to S3 for production environment."""
        if self.env != "production":
            logging.info("Not in production environment. Skipping S3 upload of static assets.")
            return

        try:
            # Get list of static files in templates directory
            static_files = [
                f for f in os.listdir(self.templates_dir) 
                if f.endswith(('.css', '.png', '.jpg', '.jpeg', '.gif'))
            ]
            
            # Upload each file to S3
            for file_name in static_files:
                local_file_path = os.path.join(self.templates_dir, file_name)
                s3_key = f"{self.static_assets_prefix}/{file_name}"
                
                logging.info(f"Uploading {file_name} to S3 bucket '{self.s3_bucket}' at '{s3_key}'")
                self.s3_client.upload_file(local_file_path, self.s3_bucket, s3_key)
                
            logging.info("All static assets uploaded to S3 successfully.")
            
        except Exception as e:
            logging.error(f"Error uploading static assets to S3: {e}")
            raise


if __name__ == "__main__":
    logging.basicConfig(
        level=logging.DEBUG, format="%(asctime)s - %(levelname)s - %(message)s"
    )
    generator = RiskMemoGenerator()
    ticker = "AAPL"
    pdf_file = generator.create_pdf(ticker)
    docx_file = generator.create_docx(ticker)
    print(f"PDF generated: {pdf_file}")
    print(f"DOCX generated: {docx_file}")
